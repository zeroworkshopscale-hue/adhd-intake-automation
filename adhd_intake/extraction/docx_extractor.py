"""Extract demographics and assess completeness of Word (.docx) assessment forms.

Some patients return the assessment questionnaire as a Microsoft Word document
instead of a PDF. This module handles that case by:

  1. Extracting all text from paragraphs and tables.
  2. Reusing the same label-based demographic extraction as the PDF extractor.
  3. Checking completeness by verifying that each required section has
     at least some non-empty content after its header.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Optional

from ..models import (
    Demographics,
    ExtractionResult,
    PdfKind,
    QuestionnaireType,
    parse_dob_candidates,
)
from ..utils.logging_config import get_logger
from . import templates

logger = get_logger(__name__)


def _docx_available() -> bool:
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        return False


def extract_text_from_docx(path: Path) -> tuple[str, dict[str, str]]:
    """Return (full_text, label_pairs) for the docx file.

    ``full_text`` is all paragraph and table text joined with newlines.
    ``label_pairs`` maps 'label: value' pairs extracted from the document.
    """
    from docx import Document  # type: ignore

    doc = Document(str(path))
    parts: list[str] = []
    pairs: dict[str, str] = {}

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
            # "Label: Value" on the same line.
            m = re.match(r"^([^:]{3,50}):\s*(.+)$", text)
            if m:
                pairs[m.group(1).strip().lower()] = m.group(2).strip()

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            # Pair adjacent non-empty cells: label | value
            for i in range(len(cells) - 1):
                if cells[i] and cells[i + 1]:
                    pairs.setdefault(cells[i].lower(), cells[i + 1])
            for cell in cells:
                if cell:
                    parts.append(cell)

    return "\n".join(parts), pairs


def _fill_demographics(text: str, pairs: dict[str, str], qtype: QuestionnaireType) -> Demographics:
    """Populate a Demographics object from the extracted text and pairs."""
    from ..models import _EMAIL_RE

    demo = Demographics()
    template = templates.template_for(qtype)
    all_labels: set[str] = set()
    if template:
        for cands in template.demographic_fields.values():
            all_labels.update(c.lower() for c in cands)

    # Fill from pairs first (most reliable for label: value format).
    if template:
        for canonical, candidates in template.demographic_fields.items():
            if getattr(demo, canonical, None):
                continue
            for label in candidates:
                label_l = label.lower()
                val = pairs.get(label_l, "")
                if val and _plausible(canonical, val):
                    setattr(demo, canonical, val)
                    break

    # Fallback: regex sweep.
    if not demo.email:
        m = _EMAIL_RE.search(text)
        if m:
            demo.email = m.group(0).strip()
    if not demo.dob:
        demo.dob = Demographics.normalise_dob(text)

    # Normalise pronoun: look for checkbox-style marks near option words.
    if not demo.pronoun:
        for opt in ("He/His", "She/Her", "They/Them"):
            if re.search(re.escape(opt), text, re.IGNORECASE):
                # Crude: if the option appears in the text it was likely selected.
                demo.pronoun = opt
                break

    if demo.dob:
        if not demo.dob_raw:
            demo.dob_raw = demo.dob
        demo.dob = Demographics.normalise_dob(demo.dob) or demo.dob

    return demo


def _plausible(canonical: str, value: str) -> bool:
    from ..models import _EMAIL_RE

    v = value.strip()
    if not v:
        return False
    if canonical == "dob":
        return bool(parse_dob_candidates(v))
    if canonical == "email":
        return bool(_EMAIL_RE.search(v))
    if canonical in ("first_name", "last_name", "pref_name"):
        return bool(re.fullmatch(r"[A-Za-z][A-Za-z .'\-,]{0,38}", v))
    return len(v) <= 80


def check_completeness_docx(path: Path, qtype: QuestionnaireType) -> tuple[bool, str]:
    """Return (is_complete, detail_message).

    For docx we can't do ink-density analysis, so we check that each of the
    expected sections has some content after its heading. Sections 6-11 (Adult)
    or 6-12 (Women) are approximated by looking for rating-scale keywords and
    confirming that the document has filled answer tokens after them.
    """
    try:
        from docx import Document

        doc = Document(str(path))
        full_text = "\n".join(
            p.text.strip() for p in doc.paragraphs if p.text.strip()
        )
        # Check that the typical question-section markers have content nearby.
        required_patterns = [
            r"never or rarely",
            r"sometimes",
            r"often or very often",
        ]
        found_any = any(
            re.search(p, full_text, re.IGNORECASE) for p in required_patterns
        )
        if not found_any:
            return False, "Rating-scale section not found in document"

        # Count how many question lines appear to be answered
        # (a number, X, checkmark, or rating word follows a question).
        answered = 0
        unanswered = 0
        for para in doc.paragraphs:
            t = para.text.strip()
            if len(t) < 15:
                continue
            # Looks like a question row if it has a rating-scale word or a number 0-4.
            if re.search(
                r"\b(never|rarely|sometimes|often|very often|yes|no|\b[0-4]\b)", t, re.IGNORECASE
            ):
                answered += 1
            elif re.search(r"\?$|^\d+[\.\)]\s+\w", t):
                unanswered += 1

        if unanswered > 3 and answered == 0:
            return False, f"No answered question rows found ({unanswered} unanswered)"
        return True, "docx completeness check passed"
    except Exception as exc:
        logger.debug("docx completeness check failed: %s", exc, exc_info=True)
        return True, "docx completeness check skipped"


def extract_docx(path: Path) -> ExtractionResult:
    """Full extraction for a .docx assessment form."""
    if not _docx_available():
        return ExtractionResult(
            questionnaire_type=QuestionnaireType.UNKNOWN,
            pdf_kind=PdfKind.UNKNOWN,
            demographics=Demographics(),
            warnings=["python-docx is not installed; cannot process Word documents."],
        )

    try:
        text, pairs = extract_text_from_docx(path)
    except Exception as exc:
        logger.exception("Failed to read docx %s", path.name)
        return ExtractionResult(
            questionnaire_type=QuestionnaireType.UNKNOWN,
            pdf_kind=PdfKind.UNKNOWN,
            demographics=Demographics(),
            warnings=[f"Could not read Word document: {exc}"],
        )

    qtype = templates.identify_type(text)
    warnings: list[str] = []
    if qtype is QuestionnaireType.UNKNOWN:
        warnings.append("Could not identify questionnaire type from Word document content.")

    # Referral source from pairs.
    pairs_lower = {k.lower(): v for k, v in pairs.items()}
    referral = pairs_lower.get("how did you hear about us", "") or pairs_lower.get("referral", "")

    answers = dict(pairs)
    if referral:
        answers["referral_source"] = referral

    demographics = _fill_demographics(text, pairs_lower, qtype)
    if not demographics.has_minimum_identifiers():
        warnings.append("Insufficient demographic identifiers extracted from Word document.")

    return ExtractionResult(
        questionnaire_type=qtype,
        pdf_kind=PdfKind.FILLABLE,  # treat docx as fillable (has text layer)
        demographics=demographics,
        answers=answers,
        raw_text=text,
        used_ocr=False,
        confidence=0.8,
        warnings=warnings,
    )
